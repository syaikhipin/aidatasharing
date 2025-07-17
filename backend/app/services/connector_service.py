"""
Enhanced Connector Service for AI Share Platform
Handles various data sources including databases, cloud storage, and document processing
"""

import logging
import json
import asyncio
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from pathlib import Path
import tempfile
import os

from sqlalchemy.orm import Session
from app.models.dataset import DatabaseConnector, Dataset, DatasetType, DatasetStatus
from app.models.organization import DataSharingLevel
from app.services.mindsdb import MindsDBService
from app.services.storage import storage_service

logger = logging.getLogger(__name__)


class ConnectorService:
    """Enhanced service for managing data connectors and document processing"""
    
    def __init__(self, db: Session):
        self.db = db
        self.mindsdb_service = MindsDBService()
        
        # Supported connector types with their configurations
        self.supported_connectors = {
            'mysql': {
                'required_config': ['host', 'port', 'database'],
                'required_credentials': ['user', 'password'],
                'optional_config': ['charset', 'ssl_disabled'],
                'default_port': 3306
            },
            'postgresql': {
                'required_config': ['host', 'port', 'database'],
                'required_credentials': ['user', 'password'],
                'optional_config': ['sslmode', 'connect_timeout'],
                'default_port': 5432
            },
            's3': {
                'required_config': ['bucket_name', 'region'],
                'required_credentials': ['aws_access_key_id', 'aws_secret_access_key'],
                'optional_config': ['endpoint_url', 'prefix'],
                'default_region': 'us-east-1'
            },
            'mongodb': {
                'required_config': ['host', 'port', 'database'],
                'required_credentials': ['username', 'password'],
                'optional_config': ['authSource', 'ssl'],
                'default_port': 27017
            },
            'snowflake': {
                'required_config': ['account', 'warehouse', 'database', 'schema'],
                'required_credentials': ['user', 'password'],
                'optional_config': ['role', 'timeout'],
                'default_schema': 'PUBLIC'
            },
            'bigquery': {
                'required_config': ['project_id', 'dataset_id'],
                'required_credentials': ['service_account_json'],
                'optional_config': ['location'],
                'default_location': 'US'
            },
            'redshift': {
                'required_config': ['host', 'port', 'database'],
                'required_credentials': ['user', 'password'],
                'optional_config': ['sslmode'],
                'default_port': 5439
            },
            'clickhouse': {
                'required_config': ['host', 'port', 'database'],
                'required_credentials': ['user', 'password'],
                'optional_config': ['secure', 'verify'],
                'default_port': 8123
            },
            'api': {
                'required_config': ['base_url', 'endpoint'],
                'required_credentials': [],
                'optional_config': ['headers', 'auth_type', 'timeout'],
                'default_timeout': 30
            },
            'file_system': {
                'required_config': ['path'],
                'required_credentials': [],
                'optional_config': ['file_pattern', 'recursive'],
                'default_recursive': False
            }
        }
        
        # Document processing capabilities
        self.document_processors = {
            'pdf': self._process_pdf_document,
            'docx': self._process_docx_document,
            'doc': self._process_doc_document,
            'txt': self._process_txt_document,
            'rtf': self._process_rtf_document,
            'odt': self._process_odt_document
        }
    
    async def create_connector_dataset(
        self,
        connector: DatabaseConnector,
        table_or_query: str,
        dataset_name: str,
        user_id: int,
        description: Optional[str] = None,
        sharing_level: DataSharingLevel = DataSharingLevel.PRIVATE
    ) -> Dict[str, Any]:
        """Create a dataset from a connector source"""
        try:
            # Validate connector is active and tested
            if not connector.is_active:
                raise Exception("Connector is not active")
            
            if connector.test_status != "success":
                raise Exception("Connector has not been successfully tested")
            
            # Create MindsDB database connection if not exists
            mindsdb_result = await self._create_mindsdb_connection(connector)
            if not mindsdb_result.get("success"):
                raise Exception(f"Failed to create MindsDB connection: {mindsdb_result.get('error')}")
            
            # Determine dataset type based on connector
            dataset_type = self._get_dataset_type_for_connector(connector.connector_type)
            
            # Create dataset record
            dataset = Dataset(
                name=dataset_name,
                description=description or f"Dataset from {connector.name} connector",
                type=dataset_type,
                status=DatasetStatus.ACTIVE,
                owner_id=user_id,
                organization_id=connector.organization_id,
                sharing_level=sharing_level,
                connector_id=connector.id,
                source_url=table_or_query,
                mindsdb_database=connector.mindsdb_database_name,
                mindsdb_table_name=table_or_query,
                allow_download=True,
                allow_api_access=True,
                allow_ai_chat=True
            )
            
            # Get schema information
            schema_info = await self._get_connector_schema(connector, table_or_query)
            if schema_info:
                dataset.schema_info = schema_info
                dataset.column_count = len(schema_info.get('columns', []))
                dataset.row_count = schema_info.get('estimated_rows', 0)
            
            self.db.add(dataset)
            self.db.commit()
            self.db.refresh(dataset)
            
            logger.info(f"✅ Created connector dataset: {dataset_name} (ID: {dataset.id})")
            
            return {
                "success": True,
                "dataset_id": dataset.id,
                "dataset_name": dataset_name,
                "connector_type": connector.connector_type,
                "mindsdb_database": connector.mindsdb_database_name,
                "source": table_or_query
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to create connector dataset: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def process_document(
        self,
        file_path: str,
        original_filename: str,
        user_id: int,
        organization_id: int,
        dataset_name: Optional[str] = None,
        description: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process document files (PDF, DOCX, etc.) and create datasets"""
        try:
            # Determine file type
            file_extension = Path(original_filename).suffix.lower().lstrip('.')
            
            if file_extension not in self.document_processors:
                raise Exception(f"Unsupported document type: {file_extension}")
            
            # Process document
            processor = self.document_processors[file_extension]
            processed_data = await processor(file_path)
            
            if not processed_data.get("success"):
                raise Exception(f"Document processing failed: {processed_data.get('error')}")
            
            # Create dataset
            dataset_name = dataset_name or Path(original_filename).stem
            
            dataset = Dataset(
                name=dataset_name,
                description=description or f"Processed document: {original_filename}",
                type=DatasetType.PDF if file_extension == 'pdf' else DatasetType.JSON,
                status=DatasetStatus.ACTIVE,
                owner_id=user_id,
                organization_id=organization_id,
                sharing_level=DataSharingLevel.PRIVATE,
                file_path=file_path,
                source_url=file_path,
                size_bytes=os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                allow_download=True,
                allow_api_access=True,
                allow_ai_chat=True,
                # Enhanced metadata
                schema_metadata=processed_data.get("metadata", {}),
                content_preview=processed_data.get("preview", ""),
                quality_metrics={
                    "document_type": file_extension,
                    "processing_method": processed_data.get("method", "unknown"),
                    "text_extraction_success": processed_data.get("text_extracted", False),
                    "page_count": processed_data.get("page_count", 0),
                    "word_count": processed_data.get("word_count", 0)
                }
            )
            
            self.db.add(dataset)
            self.db.commit()
            self.db.refresh(dataset)
            
            # Create MindsDB model for document chat
            await self._create_document_chat_model(dataset, processed_data.get("text_content", ""))
            
            logger.info(f"✅ Processed document: {original_filename} -> Dataset ID: {dataset.id}")
            
            return {
                "success": True,
                "dataset_id": dataset.id,
                "dataset_name": dataset_name,
                "document_type": file_extension,
                "processing_result": processed_data
            }
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _create_mindsdb_connection(self, connector: DatabaseConnector) -> Dict[str, Any]:
        """Create MindsDB database connection for connector"""
        try:
            connection_params = self._build_connection_string(connector)
            
            # Create database in MindsDB
            create_db_sql = f"""
            CREATE DATABASE IF NOT EXISTS {connector.mindsdb_database_name}
            WITH ENGINE = '{connector.connector_type}',
            PARAMETERS = {json.dumps(connection_params)}
            """
            
            result = self.mindsdb_service.execute_query(create_db_sql)
            
            if result.get("error"):
                return {"success": False, "error": result.get("error")}
            
            return {"success": True, "database_name": connector.mindsdb_database_name}
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _build_connection_string(self, connector: DatabaseConnector) -> Dict[str, Any]:
        """Build connection parameters for MindsDB"""
        config = connector.connection_config.copy()
        if connector.credentials:
            config.update(connector.credentials)
        
        # Connector-specific parameter mapping
        if connector.connector_type == 'mysql':
            return {
                "host": config.get("host"),
                "port": config.get("port", 3306),
                "database": config.get("database"),
                "user": config.get("user"),
                "password": config.get("password"),
                "ssl_disabled": config.get("ssl_disabled", False)
            }
        elif connector.connector_type == 'postgresql':
            return {
                "host": config.get("host"),
                "port": config.get("port", 5432),
                "database": config.get("database"),
                "user": config.get("user"),
                "password": config.get("password"),
                "sslmode": config.get("sslmode", "prefer")
            }
        elif connector.connector_type == 's3':
            return {
                "bucket": config.get("bucket_name"),
                "aws_access_key_id": config.get("aws_access_key_id"),
                "aws_secret_access_key": config.get("aws_secret_access_key"),
                "region": config.get("region", "us-east-1")
            }
        elif connector.connector_type == 'mongodb':
            return {
                "host": config.get("host"),
                "port": config.get("port", 27017),
                "database": config.get("database"),
                "username": config.get("username"),
                "password": config.get("password")
            }
        else:
            return config
    
    async def _get_connector_schema(self, connector: DatabaseConnector, table_name: str) -> Optional[Dict[str, Any]]:
        """Get schema information from connector source"""
        try:
            if connector.connector_type in ['mysql', 'postgresql']:
                # Get table schema
                schema_query = f"""
                SELECT * FROM {connector.mindsdb_database_name}.{table_name} LIMIT 0
                """
                result = self.mindsdb_service.execute_query(schema_query)
                
                if not result.get("error"):
                    # Extract column information
                    columns = []
                    if result.get("columns"):
                        columns = [{"name": col, "type": "unknown"} for col in result["columns"]]
                    
                    return {
                        "columns": columns,
                        "table_name": table_name,
                        "connector_type": connector.connector_type
                    }
            
            return None
            
        except Exception as e:
            logger.warning(f"Failed to get schema for {table_name}: {e}")
            return None
    
    def _get_dataset_type_for_connector(self, connector_type: str) -> DatasetType:
        """Map connector type to dataset type"""
        mapping = {
            'mysql': DatasetType.DATABASE,
            'postgresql': DatasetType.DATABASE,
            'mongodb': DatasetType.JSON,
            's3': DatasetType.S3_BUCKET,
            'api': DatasetType.API,
            'snowflake': DatasetType.DATABASE,
            'bigquery': DatasetType.DATABASE,
            'redshift': DatasetType.DATABASE,
            'clickhouse': DatasetType.DATABASE
        }
        return mapping.get(connector_type, DatasetType.DATABASE)
    
    # Document processing methods
    async def _process_pdf_document(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text_content = ""
            page_count = len(doc)
            
            # Extract text from all pages
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                text_content += page.get_text() + "\n"
            
            doc.close()
            
            word_count = len(text_content.split())
            
            return {
                "success": True,
                "method": "PyMuPDF",
                "text_content": text_content,
                "text_extracted": bool(text_content.strip()),
                "page_count": page_count,
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "pdf",
                    "pages": page_count,
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "PyMuPDF not installed. Install with: pip install PyMuPDF"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF processing failed: {str(e)}"
            }
    
    async def _process_docx_document(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX document using python-docx"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            word_count = len(text_content.split())
            
            return {
                "success": True,
                "method": "python-docx",
                "text_content": text_content,
                "text_extracted": bool(text_content.strip()),
                "page_count": 1,  # DOCX doesn't have explicit pages
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "python-docx not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX processing failed: {str(e)}"
            }
    
    async def _process_doc_document(self, file_path: str) -> Dict[str, Any]:
        """Process DOC document using python-docx2txt"""
        try:
            import docx2txt
            
            text_content = docx2txt.process(file_path)
            word_count = len(text_content.split())
            
            return {
                "success": True,
                "method": "docx2txt",
                "text_content": text_content,
                "text_extracted": bool(text_content.strip()),
                "page_count": 1,
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "doc",
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "docx2txt not installed. Install with: pip install docx2txt"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOC processing failed: {str(e)}"
            }
    
    async def _process_txt_document(self, file_path: str) -> Dict[str, Any]:
        """Process TXT document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            word_count = len(text_content.split())
            line_count = len(text_content.splitlines())
            
            return {
                "success": True,
                "method": "direct_read",
                "text_content": text_content,
                "text_extracted": True,
                "page_count": 1,
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "txt",
                    "lines": line_count,
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"TXT processing failed: {str(e)}"
            }
    
    async def _process_rtf_document(self, file_path: str) -> Dict[str, Any]:
        """Process RTF document using striprtf"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            
            text_content = rtf_to_text(rtf_content)
            word_count = len(text_content.split())
            
            return {
                "success": True,
                "method": "striprtf",
                "text_content": text_content,
                "text_extracted": bool(text_content.strip()),
                "page_count": 1,
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "rtf",
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "striprtf not installed. Install with: pip install striprtf"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"RTF processing failed: {str(e)}"
            }
    
    async def _process_odt_document(self, file_path: str) -> Dict[str, Any]:
        """Process ODT document using odfpy"""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(file_path)
            text_content = ""
            
            # Extract all text elements
            for element in doc.getElementsByType(text.P):
                text_content += teletype.extractText(element) + "\n"
            
            word_count = len(text_content.split())
            
            return {
                "success": True,
                "method": "odfpy",
                "text_content": text_content,
                "text_extracted": bool(text_content.strip()),
                "page_count": 1,
                "word_count": word_count,
                "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "metadata": {
                    "document_type": "odt",
                    "words": word_count,
                    "characters": len(text_content)
                }
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "odfpy not installed. Install with: pip install odfpy"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"ODT processing failed: {str(e)}"
            }
    
    async def _create_document_chat_model(self, dataset: Dataset, text_content: str) -> Dict[str, Any]:
        """Create MindsDB model for document chat"""
        try:
            model_name = f"doc_chat_{dataset.id}"
            
            # Create a simple chat model for the document
            result = self.mindsdb_service.create_gemini_model(
                model_name=model_name,
                model_type="chat",
                column_name="question"
            )
            
            if result.get("status") in ["created", "exists"]:
                # Update dataset with model information
                dataset.chat_model_name = model_name
                dataset.ai_chat_enabled = True
                dataset.chat_context = {
                    "document_content": text_content[:2000],  # First 2000 chars for context
                    "model_name": model_name,
                    "created_at": datetime.utcnow().isoformat()
                }
                self.db.commit()
                
                logger.info(f"✅ Created document chat model: {model_name}")
                return {"success": True, "model_name": model_name}
            
            return {"success": False, "error": "Model creation failed"}
            
        except Exception as e:
            logger.error(f"❌ Failed to create document chat model: {e}")
            return {"success": False, "error": str(e)}
    
    def validate_connector_config(self, connector_type: str, config: Dict[str, Any], credentials: Dict[str, Any]) -> Dict[str, Any]:
        """Validate connector configuration"""
        if connector_type not in self.supported_connectors:
            return {
                "valid": False,
                "error": f"Unsupported connector type: {connector_type}",
                "supported_types": list(self.supported_connectors.keys())
            }
        
        connector_spec = self.supported_connectors[connector_type]
        errors = []
        
        # Check required config
        for required_field in connector_spec['required_config']:
            if required_field not in config:
                errors.append(f"Missing required config field: {required_field}")
        
        # Check required credentials
        for required_field in connector_spec['required_credentials']:
            if required_field not in credentials:
                errors.append(f"Missing required credential field: {required_field}")
        
        if errors:
            return {
                "valid": False,
                "errors": errors,
                "required_config": connector_spec['required_config'],
                "required_credentials": connector_spec['required_credentials']
            }
        
        return {"valid": True}